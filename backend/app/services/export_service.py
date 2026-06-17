"""Export system.

Renders research answers and literature reviews to:
  * Markdown  (.md)
  * BibTeX    (.bib)   — citations only
  * DOCX      (.docx)  — python-docx
  * PDF       (.pdf)   — reportlab (pure-Python)

Citations are preserved across every format. The input is a normalised
``ExportDocument`` so answers, reviews, and (future) workspace reports all share
one rendering path.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any


@dataclass
class ExportDocument:
    title: str
    body_markdown: str
    citations: list[dict[str, Any]] = field(default_factory=list)
    subtitle: str | None = None
    meta: dict[str, Any] = field(default_factory=dict)


ExportFormat = str  # "markdown" | "bibtex" | "docx" | "pdf"
MIME = {
    "markdown": "text/markdown",
    "bibtex": "application/x-bibtex",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}
EXT = {"markdown": "md", "bibtex": "bib", "docx": "docx", "pdf": "pdf"}


# ── Helpers ───────────────────────────────────────────────────────────────────


def _slug(text: str) -> str:
    s = re.sub(r"[^a-zA-Z0-9]+", "-", text.lower()).strip("-")
    return (s or "noesis-export")[:60]


def _author_list(authors: Any) -> str:
    if isinstance(authors, list):
        return " and ".join(str(a) for a in authors) if authors else "Unknown"
    return str(authors) if authors else "Unknown"


def _bibtex_key(c: dict[str, Any], i: int) -> str:
    first_author = ""
    authors = c.get("authors")
    if isinstance(authors, list) and authors:
        first_author = re.sub(r"[^a-zA-Z]", "", str(authors[0]).split()[-1:][0] if str(authors[0]).split() else "")
    year = c.get("year") or "n.d."
    return f"{first_author or 'ref'}{year}{i}"


# ── Renderers ───────────────────────────────────────────────────────────────


def to_markdown(doc: ExportDocument) -> bytes:
    parts = [f"# {doc.title}"]
    if doc.subtitle:
        parts.append(f"\n*{doc.subtitle}*")
    parts.append("\n" + doc.body_markdown.strip())
    if doc.citations:
        parts.append("\n\n## References\n")
        for i, c in enumerate(doc.citations, start=1):
            line = f"{i}. {_author_list(c.get('authors'))} ({c.get('year', 'n.d.')}). "
            line += f"*{c.get('title', 'Untitled')}*."
            if c.get("venue"):
                line += f" {c['venue']}."
            if c.get("url"):
                line += f" {c['url']}"
            parts.append(line)
    return ("\n".join(parts) + "\n").encode("utf-8")


def to_bibtex(doc: ExportDocument) -> bytes:
    entries = []
    for i, c in enumerate(doc.citations, start=1):
        key = _bibtex_key(c, i)
        kind = "misc"
        if c.get("arxiv_id"):
            kind = "article"
        fields = [
            f"  title = {{{c.get('title', 'Untitled')}}}",
            f"  author = {{{_author_list(c.get('authors'))}}}",
        ]
        if c.get("year"):
            fields.append(f"  year = {{{c['year']}}}")
        if c.get("venue"):
            fields.append(f"  journal = {{{c['venue']}}}")
        if c.get("doi"):
            fields.append(f"  doi = {{{c['doi']}}}")
        if c.get("arxiv_id"):
            fields.append(f"  eprint = {{{c['arxiv_id']}}}")
        if c.get("url"):
            fields.append(f"  url = {{{c['url']}}}")
        entries.append(f"@{kind}{{{key},\n" + ",\n".join(fields) + "\n}")
    header = f"% BibTeX export — {doc.title}\n\n"
    return (header + "\n\n".join(entries) + "\n").encode("utf-8")


def to_docx(doc: ExportDocument) -> bytes:
    import docx
    from docx.shared import Pt

    document = docx.Document()
    document.add_heading(doc.title, level=0)
    if doc.subtitle:
        p = document.add_paragraph(doc.subtitle)
        p.runs[0].italic = True

    for line in doc.body_markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ")):
            document.add_paragraph(_strip_md(stripped[2:]), style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            document.add_paragraph(_strip_md(re.sub(r"^\d+\.\s", "", stripped)), style="List Number")
        else:
            document.add_paragraph(_strip_md(stripped))

    if doc.citations:
        document.add_heading("References", level=1)
        for i, c in enumerate(doc.citations, start=1):
            ref = f"[{i}] {_author_list(c.get('authors'))} ({c.get('year', 'n.d.')}). {c.get('title', 'Untitled')}."
            if c.get("venue"):
                ref += f" {c['venue']}."
            para = document.add_paragraph(ref)
            para.paragraph_format.space_after = Pt(4)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def to_pdf(doc: ExportDocument) -> bytes:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10.5,
                          leading=15, spaceAfter=7, alignment=TA_LEFT)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=15, spaceBefore=10, spaceAfter=6)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12.5, spaceBefore=8, spaceAfter=4)

    flow: list[Any] = [Paragraph(_md_inline(doc.title), styles["Title"])]
    if doc.subtitle:
        flow.append(Paragraph(f"<i>{_md_inline(doc.subtitle)}</i>", body))
    flow.append(Spacer(1, 8))

    bullets: list[Any] = []

    def _flush_bullets() -> None:
        if bullets:
            flow.append(ListFlowable(list(bullets), bulletType="bullet", leftIndent=14))
            bullets.clear()

    for line in doc.body_markdown.splitlines():
        s = line.strip()
        if not s:
            _flush_bullets()
            continue
        if s.startswith("### "):
            _flush_bullets(); flow.append(Paragraph(_md_inline(s[4:]), h2))
        elif s.startswith("## "):
            _flush_bullets(); flow.append(Paragraph(_md_inline(s[3:]), h2))
        elif s.startswith("# "):
            _flush_bullets(); flow.append(Paragraph(_md_inline(s[2:]), h1))
        elif s.startswith(("- ", "* ")):
            bullets.append(ListItem(Paragraph(_md_inline(s[2:]), body)))
        elif re.match(r"^\d+\.\s", s):
            bullets.append(ListItem(Paragraph(_md_inline(re.sub(r'^\d+\.\s', '', s)), body)))
        else:
            _flush_bullets(); flow.append(Paragraph(_md_inline(s), body))
    _flush_bullets()

    if doc.citations:
        flow.append(Paragraph("References", h1))
        for i, c in enumerate(doc.citations, start=1):
            ref = f"[{i}] {_author_list(c.get('authors'))} ({c.get('year', 'n.d.')}). {c.get('title', 'Untitled')}."
            if c.get("venue"):
                ref += f" {c['venue']}."
            flow.append(Paragraph(_md_inline(ref), body))

    pdf.build(flow)
    return buf.getvalue()


# ── Markdown → text/inline helpers ────────────────────────────────────────────


def _strip_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _md_inline(text: str) -> str:
    """Convert a subset of Markdown to reportlab's mini-HTML, escaping the rest."""
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    text = re.sub(r"`(.+?)`", r"<font face='Courier'>\1</font>", text)
    return text


RENDERERS = {
    "markdown": to_markdown,
    "bibtex": to_bibtex,
    "docx": to_docx,
    "pdf": to_pdf,
}


def render(doc: ExportDocument, fmt: ExportFormat) -> tuple[bytes, str, str]:
    """Return (bytes, mime_type, filename) for the given format."""
    fmt = fmt.lower()
    if fmt not in RENDERERS:
        raise ValueError(f"Unsupported export format: {fmt}")
    data = RENDERERS[fmt](doc)
    filename = f"{_slug(doc.title)}.{EXT[fmt]}"
    return data, MIME[fmt], filename
